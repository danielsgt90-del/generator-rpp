import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import datetime
import io
import os

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Generator RPP Deep Learning", page_icon="🎓", layout="wide")

# --- FUNGSI HELPER UNTUK WORD (DOCX) ---
def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_word_doc(data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # JUDUL
    title = doc.add_heading('RENCANA PEMBELAJARAN MENDALAM (DEEP LEARNING)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # INFO UMUM
    info_table = doc.add_table(rows=7, cols=2)
    info_data = [
        ("SATUAN PENDIDIKAN", data['satuan_pendidikan']),
        ("NAMA GURU", data['nama_guru']),
        ("MATA PELAJARAN", data['mata_pelajaran']),
        ("KELAS / SEMESTER", f"{data['kelas']} / {data['semester']}"),
        ("FASE", data['fase']),
        ("ELEMEN/MATERI POKOK", data['elemen_pokok']),
        ("ALOKASI WAKTU", data['alokasi_waktu']),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = label
        row[1].text = f": {value}"
        row[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()

    # FUNGSI INTERNAL UNTUK MEMBUAT TABEL WORD
    def add_word_table(doc, title, subtitle, headers, content_rows):
        doc.add_heading(title, level=1)
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.runs[0].italic = True
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_shading(hdr_cells[i], "D9D9D9")
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Content
        for row_data in content_rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                # Bold kolom pertama sebagai label
                if i == 0 and len(headers) == 2:
                     if row_cells[i].paragraphs[0].runs:
                        row_cells[i].paragraphs[0].runs[0].bold = True

    # TABEL 1
    t1_rows = [
        ("Peserta Didik", data['t1_peserta_didik']),
        ("Materi Pelajaran", data['t1_materi_pelajaran']),
        ("Dimensi Profil Lulusan", data['t1_profil_lulusan']),
        ("Pertanyaan Pemantik", data['t1_pertanyaan_pemantik']),
        ("Sarana & Prasarana", data['t1_sarana']),
    ]
    add_word_table(doc, "TABEL 1: IDENTIFIKASI KEBUTUHAN & KONTEKS", "(Analisis awal untuk menentukan strategi pembelajaran yang tepat)", ["Aspek", "Deskripsi Analitis"], t1_rows)

    # TABEL 2
    t2_rows = [
        ("Capaian Pembelajaran (CP)", data['t2_cp']),
        ("Tujuan Pembelajaran (TP)", data['t2_tp']),
        ("Pemahaman Bermakna", data['t2_pemahaman_bermakna']),
        ("Lintas Disiplin Ilmu", data['t2_lintas_disiplin']),
        ("Topik Pembelajaran", data['t2_topik']),
        ("Praktik Pedagogis", data['t2_pedagogis']),
        ("Kemitraan Pembelajaran", data['t2_kemitraan']),
        ("Lingkungan & Budaya Belajar", data['t2_lingkungan']),
        ("Pemanfaatan Digital", data['t2_digital']),
    ]
    add_word_table(doc, "TABEL 2: DESAIN PEMBELAJARAN", "(Peta konsep perencanaan yang menghubungkan tujuan dengan strategi)", ["Komponen", "Rumusan"], t2_rows)

    # TABEL 3
    t3_rows = [
        ("KEGIATAN AWAL\n(Opening)", data['t3_awal'], data['t3_awal_prinsip']),
        ("KEGIATAN INTI\n(Main Activities)", data['t3_inti'], data['t3_inti_prinsip']),
        ("KEGIATAN PENUTUP\n(Closing)", data['t3_penutup'], data['t3_penutup_prinsip']),
    ]
    add_word_table(doc, "TABEL 3: PENGALAMAN BELAJAR", "(Inti dari pembelajaran mendalam: Memahami, Mengaplikasi, Merefleksi)", ["Tahap", "Kegiatan Pembelajaran", "Prinsip & Strategi"], t3_rows)

    # TABEL 4
    t4_rows = [
        ("Asesmen Diagnostik\n(Awal)", data['t4_diagnostik'], data['t4_diagnostik_kriteria']),
        ("Asesmen Formatif\n(Proses)", data['t4_formatif'], data['t4_formatif_kriteria']),
        ("Asesmen Sumatif\n(Akhir)", data['t4_sumatif'], data['t4_sumatif_kriteria']),
        ("Tindak Lanjut\n(Remedial/Pengayaan)", data['t4_tindak_lanjut'], data['t4_tindak_lanjut_kriteria']),
    ]
    add_word_table(doc, "TABEL 4: ASESMEN PEMBELAJARAN", "(Penilaian yang komprehensif dan berkelanjutan)", ["Jenis Asesmen", "Teknik & Instrumen", "Kriteria/Indikator"], t4_rows)

    # TANDA TANGAN
    doc.add_paragraph()
    today = datetime.date.today()
    tt_table = doc.add_table(rows=5, cols=2)
    tt_table.rows[0].cells[0].text = "Kepala Sekolah"
    tt_table.rows[0].cells[1].text = f"{data['kota']}, {today.strftime('%d %B %Y')}"
    tt_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    tt_table.rows[4].cells[0].text = f"{data['nama_kepsek']}\nNIP. {data['nip_kepsek']}"
    tt_table.rows[4].cells[1].text = f"{data['nama_guru']}\nNIP. {data['nip_guru']}"
    tt_table.rows[4].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # LAMPIRAN
    doc.add_paragraph()
    doc.add_heading("LAMPIRAN WAJIB (Terlampir)", level=1)
    doc.add_paragraph("1. Materi Pembelajaran\n2. LKPD\n3. Rubrik Penilaian\n4. Soal Asesmen")

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- FUNGSI HELPER UNTUK PDF ---
def create_pdf_doc(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    title = Paragraph("<b>RENCANA PEMBELAJARAN MENDALAM (DEEP LEARNING)</b>", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))

    # Info Umum Table
    info_data = [
        ["SATUAN PENDIDIKAN", data['satuan_pendidikan']],
        ["NAMA GURU", data['nama_guru']],
        ["MATA PELAJARAN", data['mata_pelajaran']],
        ["KELAS / SEMESTER", f"{data['kelas']} / {data['semester']}"],
        ["FASE", data['fase']],
    ]
    info_table = Table(info_data, colWidths=[2*inch, 4.5*inch])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 20))

    # Helper for PDF Tables
    def create_pdf_table(title, data_list, col_widths):
        story.append(Paragraph(f"<b>{title}</b>", styles['Heading2']))
        table = Table(data_list, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(table)
        story.append(Spacer(1, 20))

    # Tabel 1
    t1_data = [["Aspek", "Deskripsi Analitis"]]
    t1_data.append(["Peserta Didik", data['t1_peserta_didik']])
    t1_data.append(["Materi Pelajaran", data['t1_materi_pelajaran']])
    create_pdf_table("TABEL 1: IDENTIFIKASI", t1_data, [1.5*inch, 5*inch])

    # Tabel 2
    t2_data = [["Komponen", "Rumusan"]]
    t2_data.append(["Capaian Pembelajaran", data['t2_cp']])
    t2_data.append(["Tujuan Pembelajaran", data['t2_tp']])
    create_pdf_table("TABEL 2: DESAIN PEMBELAJARAN", t2_data, [1.5*inch, 5*inch])

    # Tabel 3
    t3_data = [["Tahap", "Kegiatan", "Prinsip"]]
    t3_data.append(["Awal", data['t3_awal'], data['t3_awal_prinsip']])
    t3_data.append(["Inti", data['t3_inti'], data['t3_inti_prinsip']])
    t3_data.append(["Penutup", data['t3_penutup'], data['t3_penutup_prinsip']])
    create_pdf_table("TABEL 3: PENGALAMAN BELAJAR", t3_data, [1*inch, 3.5*inch, 2*inch])

    # Tabel 4
    t4_data = [["Jenis", "Teknik", "Kriteria"]]
    t4_data.append(["Diagnostik", data['t4_diagnostik'], data['t4_diagnostik_kriteria']])
    t4_data.append(["Formatif", data['t4_formatif'], data['t4_formatif_kriteria']])
    t4_data.append(["Sumatif", data['t4_sumatif'], data['t4_sumatif_kriteria']])
    create_pdf_table("TABEL 4: ASESMEN", t4_data, [1.2*inch, 2.5*inch, 2.8*inch])

    doc.build(story)
    buffer.seek(0)
    return buffer

# --- FUNGSI HELPER UNTUK GOOGLE DOCS (HTML) ---
def create_html_doc(data):
    # Menggunakan format HTML yang rapih yang bisa di-copy paste ke Google Docs
    html_content = f"""
    <html>
    <head>
    <style>
        body {{ font-family: 'Times New Roman', serif; margin: 40px; }}
        h1, h2 {{ text-align: center; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
        th, td {{ border: 1px solid black; padding: 8px; vertical-align: top; }}
        th {{ background-color: #d9d9d9; text-align: left; }}
        .info-label {{ font-weight: bold; width: 30%; }}
    </style>
    </head>
    <body>
        <h1>RENCANA PEMBELAJARAN MENDALAM (DEEP LEARNING)</h1>
        <table>
            <tr><td class="info-label">SATUAN PENDIDIKAN</td><td>{data['satuan_pendidikan']}</td></tr>
            <tr><td class="info-label">NAMA GURU</td><td>{data['nama_guru']}</td></tr>
            <tr><td class="info-label">MATA PELAJARAN</td><td>{data['mata_pelajaran']}</td></tr>
            <tr><td class="info-label">KELAS / SEMESTER</td><td>{data['kelas']} / {data['semester']}</td></tr>
            <tr><td class="info-label">FASE</td><td>{data['fase']}</td></tr>
        </table>

        <h2>TABEL 1: IDENTIFIKASI</h2>
        <table>
            <tr><th>Aspek</th><th>Deskripsi</th></tr>
            <tr><td>Peserta Didik</td><td>{data['t1_peserta_didik']}</td></tr>
            <tr><td>Materi Pelajaran</td><td>{data['t1_materi_pelajaran']}</td></tr>
            <tr><td>Profil Pelajar Pancasila</td><td>{data['t1_profil_lulusan']}</td></tr>
        </table>

        <h2>TABEL 2: DESAIN PEMBELAJARAN</h2>
        <table>
            <tr><th>Komponen</th><th>Rumusan</th></tr>
            <tr><td>Capaian Pembelajaran</td><td>{data['t2_cp']}</td></tr>
            <tr><td>Tujuan Pembelajaran</td><td>{data['t2_tp']}</td></tr>
            <tr><td>Pemahaman Bermakna</td><td>{data['t2_pemahaman_bermakna']}</td></tr>
        </table>
        
        <h2>TABEL 3: PENGALAMAN BELAJAR</h2>
        <table>
            <tr><th>Tahap</th><th>Kegiatan</th><th>Prinsip</th></tr>
            <tr><td>Kegiatan Awal</td><td>{data['t3_awal']}</td><td>{data['t3_awal_prinsip']}</td></tr>
            <tr><td>Kegiatan Inti</td><td>{data['t3_inti']}</td><td>{data['t3_inti_prinsip']}</td></tr>
            <tr><td>Kegiatan Penutup</td><td>{data['t3_penutup']}</td><td>{data['t3_penutup_prinsip']}</td></tr>
        </table>
    </body>
    </html>
    """
    return html_content.encode('utf-8')

# --- TAMPILAN UTAMA APLIKASI (STREAMLIT) ---

st.title("🎓 Generator RPP Deep Learning - Kurikulum Merdeka")
st.markdown("Aplikasi ini menghasilkan dokumen RPP format **Deep Learning** sesuai standar Kurikulum Merdeka SMA/SMK.")

# Inisialisasi Session State untuk menyimpan data jika halaman reload
if 'form_submitted' not in st.session_state:
    st.session_state.form_submitted = False

with st.form("rpp_form"):
    st.header("1. Informasi Umum")
    col1, col2 = st.columns(2)
    with col1:
        satuan_pendidikan = st.text_input("Satuan Pendidikan", "SMKN IV SPP-SPMA Singkawang")
        nama_guru = st.text_input("Nama Guru", "Daniel, S.Pd.K")
        mata_pelajaran = st.text_input("Mata Pelajaran", "Pendidikan Agama Kristen")
        kelas = st.selectbox("Kelas", ["X", "XI", "XII"])
        fase = st.selectbox("Fase", ["E", "F", "G"])
    with col2:
        semester = st.selectbox("Semester", ["Ganjil", "Genap"])
        elemen_pokok = st.text_input("Elemen/Materi Pokok", "Gereja dan Masyarakat Majemuk")
        alokasi_waktu = st.text_input("Alokasi Waktu", "3 X 3 JP")
        kota = st.text_input("Kota", "Singkawang")
    
    st.header("2. Identifikasi (Tabel 1)")
    t1_peserta_didik = st.text_area("Analisis Peserta Didik", "Peserta didik kelas XI berada pada fase remaja akhir...")
    t1_materi_pelajaran = st.text_area("Analisis Materi Pelajaran", "Jenis Pengetahuan: Konseptual...")
    t1_profil_lulusan = st.text_area("Dimensi Profil Pelajar Pancasila", "Beriman, Bertakwa kepada Tuhan YME...")
    t1_pertanyaan_pemantik = st.text_area("Pertanyaan Pemantik", "1. Apakah arti sesama...\n2. Bagaimana sikap gereja...")
    t1_sarana = st.text_area("Sarana & Prasarana", "Fisik: Ruang kelas. Digital: Internet...")

    st.header("3. Desain Pembelajaran (Tabel 2)")
    t2_cp = st.text_area("Capaian Pembelajaran (CP)", "Memahami perkembangan kebudayaan...")
    t2_tp = st.text_area("Tujuan Pembelajaran (TP)", "Pertemuan 1: Siswa mampu menganalisis...")
    t2_pemahaman_bermakna = st.text_area("Pemahaman Bermakna", "Memahami bahwa gereja hadir...")
    t2_lintas_disiplin = st.text_input("Lintas Disiplin Ilmu", "PPKn, Sosiologi")
    t2_topik = st.text_input("Topik Pembelajaran", "Gereja sebagai Salib Bagi Dunia")
    t2_pedagogis = st.text_area("Praktik Pedagogis", "Model: Project Based Learning...")
    t2_kemitraan = st.text_area("Kemitraan Pembelajaran", "Orang tua, Tokoh Agama...")
    t2_lingkungan = st.text_area("Lingkungan & Budaya Belajar", "Fisik: Kelompok. Sosial: Safe space...")
    t2_digital = st.text_area("Pemanfaatan Digital", "Canva, Youtube, Google Forms...")

    st.header("4. Pengalaman Belajar (Tabel 3)")
    t3_awal = st.text_area("Kegiatan Awal", "1. Salam doa.\n2. Apersepsi...")
    t3_awal_prinsip = st.text_input("Prinsip Awal", "Menggembirakan & Bermakna")
    t3_inti = st.text_area("Kegiatan Inti", "A. Memahami: Studi Alkitab...\nB. Mengaplikasi: Diskusi...")
    t3_inti_prinsip = st.text_input("Prinsip Inti", "Berkesadaran & Bermakna")
    t3_penutup = st.text_area("Kegiatan Penutup", "1. Rangkuman.\n2. Refleksi diri...")
    t3_penutup_prinsip = st.text_input("Prinsip Penutup", "Berkesadaran & Menggembirakan")

    st.header("5. Asesmen (Tabel 4)")
    col_a1, col_a2 = st.columns(2)
    with col_a1:
        t4_diagnostik = st.text_area("Asesmen Diagnostik (Teknik)", "Kuis diagnostik awal.")
        t4_formatif = st.text_area("Asesmen Formatif (Teknik)", "Observasi keaktifan diskusi.")
        t4_sumatif = st.text_area("Asesmen Sumatif (Teknik)", "Proyek Poster/Video Toleransi.")
    with col_a2:
        t4_diagnostik_kriteria = st.text_area("Kriteria Diagnostik", "Mengukur pengetahuan prasyarat.")
        t4_formatif_kriteria = st.text_area("Kriteria Formatif", "Kolaborasi dan kualitas argumentasi.")
        t4_sumatif_kriteria = st.text_area("Kriteria Sumatif", "Pencapaian KKO sesuai TP.")
    t4_tindak_lanjut = st.text_area("Tindak Lanjut", "Remedial: Bimbingan individu...")
    t4_tindak_lanjut_kriteria = st.text_area("Kriteria Tindak Lanjut", "Sesuai hasil asesmen.")

    st.header("6. Penandatangan")
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        nama_kepsek = st.text_input("Nama Kepala Sekolah", "Nama Kepala Sekolah")
        nip_kepsek = st.text_input("NIP Kepala Sekolah", ".......................")
    with col_s2:
        nip_guru = st.text_input("NIP Guru", ".......................")

    submitted = st.form_submit_button("🚀 Generate RPP")

# --- LOGIKA EKSEKUSI ---
if submitted:
    # Kumpulkan semua data ke dalam dictionary
    data = {
        'satuan_pendidikan': satuan_pendidikan, 'nama_guru': nama_guru, 'mata_pelajaran': mata_pelajaran,
        'kelas': kelas, 'semester': semester, 'fase': fase, 'elemen_pokok': elemen_pokok,
        'alokasi_waktu': alokasi_waktu, 'kota': kota,
        't1_peserta_didik': t1_peserta_didik, 't1_materi_pelajaran': t1_materi_pelajaran,
        't1_profil_lulusan': t1_profil_lulusan, 't1_pertanyaan_pemantik': t1_pertanyaan_pemantik,
        't1_sarana': t1_sarana,
        't2_cp': t2_cp, 't2_tp': t2_tp, 't2_pemahaman_bermakna': t2_pemahaman_bermakna,
        't2_lintas_disiplin': t2_lintas_disiplin, 't2_topik': t2_topik, 't2_pedagogis': t2_pedagogis,
        't2_kemitraan': t2_kemitraan, 't2_lingkungan': t2_lingkungan, 't2_digital': t2_digital,
        't3_awal': t3_awal, 't3_awal_prinsip': t3_awal_prinsip, 't3_inti': t3_inti,
        't3_inti_prinsip': t3_inti_prinsip, 't3_penutup': t3_penutup, 't3_penutup_prinsip': t3_penutup_prinsip,
        't4_diagnostik': t4_diagnostik, 't4_diagnostik_kriteria': t4_diagnostik_kriteria,
        't4_formatif': t4_formatif, 't4_formatif_kriteria': t4_formatif_kriteria,
        't4_sumatif': t4_sumatif, 't4_sumatif_kriteria': t4_sumatif_kriteria,
        't4_tindak_lanjut': t4_tindak_lanjut, 't4_tindak_lanjut_kriteria': t4_tindak_lanjut_kriteria,
        'nama_kepsek': nama_kepsek, 'nip_kepsek': nip_kepsek, 'nip_guru': nip_guru
    }
    
    st.success("✅ Dokumen Berhasil Dibuat! Silakan pilih format unduhan di bawah ini.")
    
    col_btn1, col_btn2, col_btn3 = st.columns(3)
    
    # Generate Word
    with col_btn1:
        word_buffer = create_word_doc(data)
        st.download_button(
            label="📥 Download Ms. Word (.docx)",
            data=word_buffer,
            file_name=f"RPP_{data['mata_pelajaran']}_{data['nama_guru'].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Generate PDF
    with col_btn2:
        pdf_buffer = create_pdf_doc(data)
        st.download_button(
            label="📥 Download PDF (.pdf)",
            data=pdf_buffer,
            file_name=f"RPP_{data['mata_pelajaran']}_{data['nama_guru'].replace(' ', '_')}.pdf",
            mime="application/pdf"
        )

    # Generate HTML (Google Docs)
    with col_btn3:
        html_content = create_html_doc(data)
        st.download_button(
            label="📥 Download Google Docs (.html)",
            data=html_content,
            file_name=f"RPP_{data['mata_pelajaran']}_{data['nama_guru'].replace(' ', '_')}.html",
            mime="text/html"
        )
        st.caption("💡 *Tips: Buka file HTML ini di browser, copy semua (Ctrl+A), lalu paste ke dalam Google Docs.*")
